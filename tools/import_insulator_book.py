import argparse
import hashlib
import html
import json
import mimetypes
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


TEXT_EXTENSIONS = {".docx", ".xlsx", ".pdf"}
ASSET_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".jpg", ".jpeg", ".png", ".gif", ".webp"}
MAX_BUNDLED_ASSET_MB = 10


TOPIC_DEFINITIONS = [
    {
        "id": "field-notebook",
        "title": "Field Notebook",
        "description": "Short Snowman's Notes lessons, field references, sketches, and working notebook sections.",
        "keywords": ["snowmans notes", "snowman", "notebook", "journal", "jon"],
    },
    {
        "id": "pattern-development",
        "title": "Pattern Development",
        "description": "Short layout lessons for parallel line development, tees, bends, gores, and fitting geometry.",
        "keywords": ["development", "tee", "gore", "pattern", "map", "dimension", "layout", "bend", "body"],
    },
    {
        "id": "reference-charts",
        "title": "Reference Charts",
        "description": "Quick lookup charts, size tables, cards, formulas, and field reference topics.",
        "keywords": ["chart", "reference", "size", "section", "field_reference", "field reference", "conduit", "cards"],
    },
    {
        "id": "book-supplements",
        "title": "Book Supplements",
        "description": "Supplemental lessons, worksheets, local book sections, measuring guides, and supporting notes.",
        "keywords": ["supplement", "local 7", "type me", "metal", "band", "measuring", "jons notes"],
    },
    {
        "id": "insulator-tools",
        "title": "Insulator Tools",
        "description": "Tool-folder lessons converted into readable text-and-image pages instead of downloads.",
        "keywords": ["insulator tools"],
    },
    {
        "id": "firestopping-safety",
        "title": "Firestopping & Safety",
        "description": "Short safety references for firestopping, asbestos, hazards, controls, and inspections.",
        "keywords": ["fire", "firestop", "asbestos", "hazard", "safety", "controls"],
    },
    {
        "id": "standards-specs",
        "title": "Standards & Specs",
        "description": "Standards, specs, handbooks, working rules, and official guidance as reference topics.",
        "keywords": ["standard", "spec", "handbook", "rules", "cbr", "commercial", "industrial"],
    },
    {
        "id": "sketches-images",
        "title": "Sketches & Images",
        "description": "Standalone visual references, diagrams, sketches, and extracted page images.",
        "keywords": ["sketch", "image", "jpg", "png", "t-master", "hazards and controls"],
    },
    {
        "id": "source-inventory",
        "title": "Complete Source Inventory",
        "description": "Every source file found in the provided folder, kept as the archive behind the lessons.",
        "keywords": [],
        "include_all": True,
    },
]


def slugify(value):
    value = value.lower().replace("&", " and ")
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "file"


def sha1_short(text):
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:10]


def clean_text(value):
    value = re.sub(r"\s+", " ", value or "").strip()
    return value


def classify(path):
    name = path.stem.lower()
    parent = " ".join(p.name.lower() for p in path.parents)
    haystack = f"{name} {parent}"

    if "insulator tools" in haystack:
        return "Insulator Tools"
    if "snowmans notes half size" in name:
        return "Notebook"
    if "reference" in haystack or "chart" in haystack or "size chart" in haystack:
        return "Reference Charts"
    if "supplement" in haystack:
        return "Book Supplements"
    if "layout dimension" in haystack or "development" in haystack or "tee" in haystack or "gore" in haystack:
        return "Layouts & Patterns"
    if "fire" in haystack or "asbestos" in haystack or "hazard" in haystack:
        return "Safety & Firestopping"
    if "standard" in haystack or "spec" in haystack or "rules" in haystack or "handbook" in haystack:
        return "Standards & Handbooks"
    if "tool" in haystack:
        return "Insulator Tools"
    if path.suffix.lower() in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
        return "Sketches & Images"
    return "Library"


def assign_topics(path, category):
    relative = str(path).lower().replace("\\", "/")
    topics = []
    for topic in TOPIC_DEFINITIONS:
        if topic.get("include_all"):
            continue
        if any(keyword in relative for keyword in topic["keywords"]):
            topics.append(topic["id"])

    category_map = {
        "Notebook": "field-notebook",
        "Layouts & Patterns": "pattern-development",
        "Reference Charts": "reference-charts",
        "Book Supplements": "book-supplements",
        "Insulator Tools": "insulator-tools",
        "Safety & Firestopping": "firestopping-safety",
        "Standards & Handbooks": "standards-specs",
        "Sketches & Images": "sketches-images",
    }
    mapped = category_map.get(category)
    if mapped and mapped not in topics:
        topics.append(mapped)
    if not topics:
        topics.append("source-inventory")
    topics.append("source-inventory")
    return list(dict.fromkeys(topics))


def extract_docx(path, image_out, rel_image_root):
    paragraphs = []
    tables = []
    image_count = 0
    images = []
    doc = Document(path)

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            style = paragraph.style.name if paragraph.style else ""
            paragraphs.append({"text": text, "style": style})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    try:
        with ZipFile(path) as zf:
            for member in zf.namelist():
                if member.startswith("word/media/"):
                    ext = Path(member).suffix.lower()
                    if ext in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                        image_count += 1
                        image_out.mkdir(parents=True, exist_ok=True)
                        out_name = f"{slugify(path.stem)}-{image_count:03d}{ext}"
                        out = image_out / out_name
                        with zf.open(member) as src, out.open("wb") as dst:
                            shutil.copyfileobj(src, dst)
                        images.append(f"{rel_image_root}/{out_name}")
    except Exception:
        pass

    table_text = [" | ".join(cell for cell in row if cell) for table in tables for row in table[:8]]
    text = "\n".join([p["text"] for p in paragraphs] + table_text)
    return {
        "text": text,
        "paragraphs": paragraphs[:250],
        "tables": tables[:30],
        "image_count": image_count,
        "images": images,
        "extraction": "docx",
    }


def extract_xlsx(path):
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheets = []
    text_parts = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(max_row=250, values_only=True):
            values = [clean_text(str(value)) if value is not None else "" for value in row]
            if any(values):
                rows.append(values)
                text_parts.append(" | ".join(v for v in values if v))
        sheets.append({"name": sheet.title, "rows": rows[:120]})
    return {
        "text": "\n".join(text_parts),
        "sheets": sheets,
        "extraction": "xlsx",
    }


def extract_pdf(path, image_out, rel_image_root, extract_images=False):
    pages = 0
    text_parts = []
    page_texts = []
    images = []
    status = "pdf-text"
    try:
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        for page_index, page in enumerate(reader.pages[:60], start=1):
            raw_text = page.extract_text() or ""
            lines = [clean_text(line) for line in raw_text.splitlines()]
            page_text = "\n".join(line for line in lines if line)
            if page_text:
                text_parts.append(page_text)
                page_texts.append({"page": page_index, "text": page_text})
            if extract_images:
                for image_index, image in enumerate(getattr(page, "images", [])[:6], start=1):
                    if len(images) >= 80:
                        break
                    ext = Path(getattr(image, "name", "")).suffix.lower()
                    if ext not in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
                        ext = ".jpg"
                    image_out.mkdir(parents=True, exist_ok=True)
                    out_name = f"{slugify(path.stem)}-p{page_index:03d}-{image_index:02d}{ext}"
                    out = image_out / out_name
                    out.write_bytes(image.data)
                    images.append(f"{rel_image_root}/{out_name}")
        if not text_parts:
            status = "pdf-needs-ocr"
    except Exception as exc:
        status = f"pdf-error: {exc.__class__.__name__}"
    return {
        "text": "\n".join(text_parts),
        "pdfPages": page_texts,
        "images": images,
        "pages": pages,
        "extraction": status,
    }


def summarize(text, fallback):
    text = clean_text(text)
    if not text:
        return fallback
    sentences = re.split(r"(?<=[.!?])\s+", text)
    summary = " ".join(sentences[:2])
    return summary[:360].strip()


def copy_asset(path, public_dir, rel_root):
    size_mb = path.stat().st_size / (1024 * 1024)
    if path.suffix.lower() not in ASSET_EXTENSIONS:
        return None, "unsupported"
    if size_mb > MAX_BUNDLED_ASSET_MB:
        return None, "oversized"
    target_name = f"{slugify(path.stem)}-{sha1_short(str(path))}{path.suffix.lower()}"
    target = public_dir / target_name
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, target)
    return f"{rel_root}/{target_name}", "bundled"


def empty_directory(path):
    path.mkdir(parents=True, exist_ok=True)
    for child in path.rglob("*"):
        if child.is_file():
            child.unlink()


def file_entry(path, source_root, asset_dir, rel_asset_root):
    relative = str(path.relative_to(source_root)).replace("\\", "/")
    is_tool_resource = relative.lower().startswith("insulator tools/")
    stat = path.stat()
    ext = path.suffix.lower()
    extracted = {"text": "", "extraction": "not-parsed"}

    try:
        if ext == ".docx":
            extracted = extract_docx(path, asset_dir / "docx-images", "assets/docx-images")
        elif ext == ".xlsx":
            extracted = extract_xlsx(path)
        elif ext == ".pdf":
            extracted = extract_pdf(path, asset_dir / "pdf-images", "assets/pdf-images", extract_images=is_tool_resource)
    except Exception as exc:
        extracted = {"text": "", "extraction": f"error: {exc.__class__.__name__}"}

    if is_tool_resource:
        asset_url, asset_status = None, "text-resource"
    else:
        asset_url, asset_status = copy_asset(path, asset_dir / "files", f"{rel_asset_root}/files")
    summary = summarize(extracted.get("text", ""), f"{path.stem} ({ext.lstrip('.').upper()})")
    title = path.stem.replace("_", " ").replace("  ", " ").strip()
    category = classify(path)
    topics = assign_topics(path.relative_to(source_root), category)

    return {
        "id": slugify(f"{relative}-{sha1_short(relative)}"),
        "title": title,
        "category": category,
        "topics": topics,
        "extension": ext.lstrip(".").upper(),
        "relativePath": relative,
        "sourcePath": str(path),
        "sizeBytes": stat.st_size,
        "sizeMB": round(stat.st_size / (1024 * 1024), 2),
        "modified": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
        "summary": summary,
        "assetUrl": asset_url,
        "assetStatus": asset_status,
        "downloadAllowed": not is_tool_resource,
        "resourceMode": "text-resource" if is_tool_resource else "file-reference",
        "extraction": extracted.get("extraction", "not-parsed"),
        "pages": extracted.get("pages"),
        "imageCount": extracted.get("image_count", 0),
        "images": extracted.get("images", []),
        "thumbnailUrl": asset_url if ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"} else (extracted.get("images", [None]) or [None])[0],
        "pdfPages": extracted.get("pdfPages", []),
        "paragraphs": extracted.get("paragraphs", []),
        "tables": extracted.get("tables", []),
        "sheets": extracted.get("sheets", []),
        "searchText": clean_text(" ".join([title, category, relative, extracted.get("text", "")]))[:60000],
    }


def xlsx_sheet_entries(entry):
    if entry.get("extension") != "XLSX" or not entry.get("sheets"):
        return []

    sheet_entries = []
    for sheet in entry["sheets"]:
        sheet_name = sheet["name"]
        sheet_text = "\n".join(
            " | ".join(cell for cell in row if cell)
            for row in sheet.get("rows", [])
        )
        title = sheet_name.strip() or f"{entry['title']} sheet"
        relative = f"{entry['relativePath']}#{sheet_name}"
        sheet_entry = {
            **entry,
            "id": slugify(f"{entry['relativePath']}-{sheet_name}-{sha1_short(relative)}"),
            "title": title,
            "relativePath": relative,
            "sourcePath": f"{entry['sourcePath']}#{sheet_name}",
            "summary": summarize(sheet_text, f"{entry['title']} workbook tab"),
            "resourceMode": "xlsx-sheet",
            "sheets": [sheet],
            "searchText": clean_text(" ".join([title, entry["category"], relative, sheet_text]))[:60000],
        }
        sheet_entries.append(sheet_entry)
    return sheet_entries


def expand_workbook_sheet_entries(entries):
    expanded = []
    for entry in entries:
        expanded.append(entry)
        expanded.extend(xlsx_sheet_entries(entry))
    return expanded


def resource_kind(entry):
    haystack = " ".join([
        entry.get("title", ""),
        entry.get("category", ""),
        entry.get("relativePath", ""),
    ]).lower()
    if entry.get("resourceMode") == "xlsx-sheet" or re.search(r"chart|table|card|size|flange|fraction|decimal|reference", haystack):
        return "Quick reference"
    if re.search(r"lesson|unit|development|tee|gore|layout|pattern|method|formula|section|handbook|notebook", haystack):
        return "Short lesson"
    if entry.get("thumbnailUrl") or entry.get("extension") in {"JPG", "JPEG", "PNG", "GIF", "WEBP"}:
        return "Image reference"
    return "Source record"


def entry_card(entry, prefix="../"):
    status = ""
    if entry.get("assetStatus") == "text-resource":
        status = "Text resource"
    elif entry.get("assetStatus") == "oversized":
        status = "Listed only: oversized for static hosting"
    elif entry.get("assetUrl"):
        status = "File included"
    elif entry.get("extraction") == "pdf-needs-ocr":
        status = "Needs OCR"
    else:
        status = entry.get("extraction", "")

    thumb = ""
    if entry.get("thumbnailUrl"):
        thumb = f"<img class=\"card-thumb\" src=\"{prefix}{html.escape(entry['thumbnailUrl'])}\" alt=\"\">"

    actions = [f"<a class=\"button detail-link\" href=\"{prefix}{html.escape(entry['pageUrl'])}\">Read</a>"]
    if entry.get("assetUrl") and entry.get("downloadAllowed", True):
        actions.append(f"<a class=\"button ghost asset-link\" href=\"{prefix}{html.escape(entry['assetUrl'])}\">Open file</a>")

    return f"""
<article class=\"resource-card\">
  {thumb}
  <div class=\"card-top\"><span class=\"pill\">{html.escape(resource_kind(entry))}</span><span class=\"file-size\">{entry['sizeMB']} MB</span></div>
  <h3>{html.escape(entry['title'])}</h3>
  <p class=\"summary\">{html.escape(entry['summary'])}</p>
  <div class=\"card-actions\">{''.join(actions)}</div>
  <p class=\"status\">{html.escape(status)}</p>
</article>
"""


def render_topic_page(topic, entries, out_dir):
    topic_entries = entries if topic.get("include_all") else [entry for entry in entries if topic["id"] in entry.get("topics", [])]
    topic_entries = sorted(topic_entries, key=lambda entry: (entry["category"], entry["title"].lower(), entry["relativePath"].lower()))
    readable = [
        entry for entry in topic_entries
        if entry.get("paragraphs") or entry.get("pdfPages") or entry.get("sheets") or entry.get("images")
    ]
    highlights = []
    for entry in readable[:8]:
        highlights.append(
            f"<section class=\"topic-excerpt\"><h2>{html.escape(entry['title'])}</h2>"
            f"<p>{html.escape(entry['summary'])}</p>"
            f"<p><a class=\"back-link\" href=\"../{html.escape(entry['pageUrl'])}\">Open full resource</a></p></section>"
        )
    if not highlights:
        highlights.append("<section class=\"topic-excerpt\"><h2>Source Listing</h2><p>This topic is represented through the source list below.</p></section>")

    cards = "".join(entry_card(entry) for entry in topic_entries)
    topic_dir = out_dir / "topics"
    topic_dir.mkdir(parents=True, exist_ok=True)
    page = f"""<!doctype html>
<html lang=\"en\">
<head>
  <meta charset=\"utf-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
  <title>{html.escape(topic['title'])}</title>
  <link rel=\"stylesheet\" href=\"../styles.css\">
</head>
<body class=\"detail-page\">
  <main class=\"topic-page\">
    <section class=\"sheet detail topic-hero\">
      <a href=\"../index.html\" class=\"back-link\">Back to field notebook</a>
      <p class=\"eyebrow\">Lessons & References</p>
      <h1>{html.escape(topic['title'])}</h1>
      <p>{html.escape(topic['description'])}</p>
      <p class=\"meta\">{len(topic_entries)} lessons and references included</p>
    </section>
    <section class=\"topic-layout\">
      <div class=\"topic-main\">{''.join(highlights)}</div>
      <aside class=\"topic-source-list sheet\">
        <h2>Lessons & Reference Cards</h2>
        <div class=\"card-grid compact\">{cards}</div>
      </aside>
    </section>
  </main>
</body>
</html>
"""
    (topic_dir / f"{topic['id']}.html").write_text(page, encoding="utf-8")
    return {
        "id": topic["id"],
        "title": topic["title"],
        "description": topic["description"],
        "url": f"topics/{topic['id']}.html",
        "count": len(topic_entries),
    }


def build_topic_pages(entries, out_dir):
    return [render_topic_page(topic, entries, out_dir) for topic in TOPIC_DEFINITIONS]


def build_html_pages(entries, out_dir):
    pages_dir = out_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    for entry in entries:
        body = [f"<h1>{html.escape(entry['title'])}</h1>"]
        body.append(f"<p class=\"meta\">{html.escape(entry['category'])} · {entry['extension']} · {entry['sizeMB']} MB</p>")
        body.append(f"<p>{html.escape(entry['summary'])}</p>")
        if entry.get("assetUrl") and entry.get("downloadAllowed", True):
            body.append(f"<p><a class=\"button\" href=\"../{html.escape(entry['assetUrl'])}\">Open file</a></p>")
        elif entry.get("assetStatus") == "text-resource":
            body.append("<p class=\"notice\">This tool has been converted to a text-based web resource. The source file is not offered as a download here.</p>")
        elif entry["assetStatus"] == "oversized":
            body.append("<p class=\"notice\">This source is too large to bundle for free static hosting. Keep it in cloud storage and add its public link in the catalog when ready.</p>")
        if entry.get("thumbnailUrl"):
            body.append(f"<figure class=\"lead-image\"><img src=\"../{html.escape(entry['thumbnailUrl'])}\" alt=\"{html.escape(entry['title'])}\"></figure>")
        for paragraph in entry.get("paragraphs", [])[:120]:
            text = html.escape(paragraph["text"])
            style = paragraph.get("style", "")
            if "Heading 1" in style or "Title" in style:
                body.append(f"<h2>{text}</h2>")
            elif "Heading" in style:
                body.append(f"<h3>{text}</h3>")
            else:
                body.append(f"<p>{text}</p>")
        for table in entry.get("tables", [])[:10]:
            body.append("<div class=\"table-wrap\"><table>")
            for idx, row in enumerate(table[:40]):
                tag = "th" if idx == 0 else "td"
                cells = "".join(f"<{tag}>{html.escape(cell)}</{tag}>" for cell in row[:8])
                body.append(f"<tr>{cells}</tr>")
            body.append("</table></div>")
        for pdf_page in entry.get("pdfPages", [])[:60]:
            body.append(f"<section class=\"text-page\"><h2>Page {pdf_page['page']}</h2>")
            for paragraph in pdf_page["text"].split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    body.append(f"<p>{html.escape(paragraph)}</p>")
            body.append("</section>")
        if entry.get("images"):
            body.append("<h2>Notebook Images</h2><div class=\"image-grid\">")
            for image_url in entry["images"][:80]:
                body.append(f"<a href=\"../{html.escape(image_url)}\"><img src=\"../{html.escape(image_url)}\" alt=\"Notebook image\"></a>")
            body.append("</div>")
        for sheet in entry.get("sheets", [])[:8]:
            body.append(f"<h2>{html.escape(sheet['name'])}</h2><div class=\"table-wrap\"><table>")
            for idx, row in enumerate(sheet["rows"][:80]):
                tag = "th" if idx == 0 else "td"
                cells = "".join(f"<{tag}>{html.escape(cell)}</{tag}>" for cell in row[:10])
                body.append(f"<tr>{cells}</tr>")
            body.append("</table></div>")

        page = f"""<!doctype html>
<html lang=\"en\">
<head>
  <meta charset=\"utf-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
  <title>{html.escape(entry['title'])}</title>
  <link rel=\"stylesheet\" href=\"../styles.css\">
</head>
<body class=\"detail-page\">
  <main class=\"sheet detail\">
    <a href=\"../index.html\" class=\"back-link\">Back to field notebook</a>
    {''.join(body)}
  </main>
</body>
</html>
"""
        (pages_dir / f"{entry['id']}.html").write_text(page, encoding="utf-8")
        entry["pageUrl"] = f"pages/{entry['id']}.html"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--out", default=".")
    args = parser.parse_args()

    source_root = Path(args.source).resolve()
    out_dir = Path(args.out).resolve()
    data_dir = out_dir / "data"
    asset_dir = out_dir / "assets"
    data_dir.mkdir(parents=True, exist_ok=True)
    asset_dir.mkdir(parents=True, exist_ok=True)

    for generated in [data_dir, asset_dir, out_dir / "pages", out_dir / "topics"]:
        empty_directory(generated)

    files = [
        p for p in source_root.rglob("*")
        if p.is_file() and p.name.lower() != "desktop.ini" and not p.name.startswith("~$")
    ]
    source_file_count = len(files)
    entries = [file_entry(p, source_root, asset_dir, "assets") for p in sorted(files)]
    entries = expand_workbook_sheet_entries(entries)
    build_html_pages(entries, out_dir)
    topics = build_topic_pages(entries, out_dir)

    categories = sorted({entry["category"] for entry in entries})
    notebook = next((e for e in entries if e["title"].lower() == "snowmans notes half size" and e["extension"] == "DOCX"), None)
    catalog = {
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "sourceRoot": str(source_root),
        "sourceFileCount": source_file_count,
        "resourceCount": len(entries),
        "maxBundledAssetMB": MAX_BUNDLED_ASSET_MB,
        "categories": categories,
        "topics": topics,
        "notebookId": notebook["id"] if notebook else None,
        "entries": entries,
    }
    (data_dir / "catalog.json").write_text(json.dumps(catalog, indent=2), encoding="utf-8")
    print(f"Imported {source_file_count} source files into {len(entries)} site resources in {out_dir}")


if __name__ == "__main__":
    main()
